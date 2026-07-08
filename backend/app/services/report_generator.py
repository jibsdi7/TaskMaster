"""
Execution Report Generator
Builds a .docx document from a WorkflowRun + its WorkflowLog rows.
Each log entry that has an associated screenshot embeds the image inline.
"""
from __future__ import annotations

import io
import os
from datetime import datetime
from zoneinfo import ZoneInfo
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from app.db import models


class ReportGenerator:
    """Generate execution reports in DOCX format."""

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        """Shade a table cell with a hex fill colour."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Insert a thin paragraph border acting as a horizontal rule."""
        paragraph = doc.add_paragraph()
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        paragraph_border.append(bottom)
        paragraph_properties.append(paragraph_border)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    _TZ = ZoneInfo("America/New_York")

    def _format_datetime(self, value: datetime | None) -> str:
        if value is None:
            return "—"
        if value.tzinfo is None:
            value = value.replace(tzinfo=ZoneInfo("UTC"))
        return value.astimezone(self._TZ).strftime("%d %b %Y  %H:%M:%S %Z")

    def _level_color(self, level: str) -> RGBColor:
        return {
            "INFO": RGBColor(0x3B, 0x82, 0xD4),
            "WARNING": RGBColor(0xF6, 0xAD, 0x55),
            "ERROR": RGBColor(0xF5, 0x65, 0x65),
        }.get(level.upper(), RGBColor(0x88, 0x88, 0x88))

    def generate_execution_report(
        self,
        workflow: "models.Workflow",
        run: "models.WorkflowRun",
        logs: list["models.WorkflowLog"],
    ) -> io.BytesIO:
        document = Document()

        for section in document.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run("Execution Report")
        title_run.bold = True
        title_run.font.size = Pt(26)
        title_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
        title_paragraph.paragraph_format.space_after = Pt(4)

        subtitle_paragraph = document.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(workflow.name)
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
        subtitle_paragraph.paragraph_format.space_after = Pt(16)

        self._add_horizontal_rule(document)
        document.add_paragraph()

        table = document.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        status_colors = {
            "completed": "48BB78",
            "failed": "F56565",
            "active": "7B96F9",
            "draft": "F6AD55",
        }

        status_str = str(run.status.value if hasattr(run.status, "value") else run.status)
        status_color = status_colors.get(status_str.lower(), "888888")
        rows_data = [
            ("Run ID", run.run_id),
            ("Status", status_str.upper()),
            ("Started", self._format_datetime(run.started_at)),
            ("Completed", self._format_datetime(run.completed_at)),
            ("Duration", f"{run.duration_seconds:.2f}s" if run.duration_seconds is not None else "—"),
        ]

        for index, (key, value) in enumerate(rows_data):
            row = table.rows[index]
            key_cell = row.cells[0]
            key_cell.text = key
            key_run = key_cell.paragraphs[0].runs[0]
            key_run.bold = True
            key_run.font.size = Pt(10)
            key_run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)
            self._set_cell_bg(key_cell, "F7F8FA")

            value_cell = row.cells[1]
            value_cell.text = value
            value_run = value_cell.paragraphs[0].runs[0]
            value_run.font.size = Pt(10)
            if key == "Status":
                value_run.bold = True
                value_run.font.color.rgb = RGBColor(
                    int(status_color[0:2], 16),
                    int(status_color[2:4], 16),
                    int(status_color[4:6], 16),
                )

        if run.error_message:
            error_paragraph = document.add_paragraph()
            error_paragraph.paragraph_format.space_before = Pt(10)
            error_label = error_paragraph.add_run("Error:  ")
            error_label.bold = True
            error_label.font.color.rgb = RGBColor(0xF5, 0x65, 0x65)
            error_paragraph.add_run(run.error_message).font.color.rgb = RGBColor(0xF5, 0x65, 0x65)

        document.add_paragraph()
        self._add_horizontal_rule(document)

        steps_heading = document.add_paragraph()
        heading_run = steps_heading.add_run("Execution Steps")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
        steps_heading.paragraph_format.space_before = Pt(14)
        steps_heading.paragraph_format.space_after = Pt(6)

        step_number = 0
        for log in logs:
            level = (log.level or "INFO").upper()
            if level == "DEBUG":
                continue
            is_action = level == "INFO" and log.node_id and "executed successfully" in (log.message or "")
            if is_action:
                step_number += 1

            step_paragraph = document.add_paragraph()
            step_paragraph.paragraph_format.space_before = Pt(10)
            step_paragraph.paragraph_format.space_after = Pt(2)

            if is_action:
                step_run = step_paragraph.add_run(f"Step {step_number}  ")
                step_run.bold = True
                step_run.font.size = Pt(11)
                step_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)

            level_run = step_paragraph.add_run(f"[{level}]  ")
            level_run.bold = True
            level_run.font.size = Pt(10)
            level_run.font.color.rgb = self._level_color(level)

            message_run = step_paragraph.add_run(log.message or "")
            message_run.font.size = Pt(10)
            message_run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)

            meta_paragraph = document.add_paragraph()
            meta_paragraph.paragraph_format.space_before = Pt(0)
            meta_paragraph.paragraph_format.space_after = Pt(2)
            meta_parts = []
            if log.node_id:
                meta_parts.append(f"node: {log.node_id}")
            if log.created_at:
                ts = log.created_at
                if ts.tzinfo is None:
                    ts = ts.replace(tzinfo=ZoneInfo("UTC"))
                meta_parts.append(ts.astimezone(self._TZ).strftime("%H:%M:%S %Z"))
            meta_run = meta_paragraph.add_run("  ".join(meta_parts))
            meta_run.font.size = Pt(8)
            meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            meta_run.italic = True

            if log.screenshot_path:
                absolute_path = os.path.abspath(log.screenshot_path)
                if not os.path.isfile(absolute_path):
                    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
                    absolute_path = os.path.join(backend_dir, log.screenshot_path.replace("\\", os.sep))

                if os.path.isfile(absolute_path):
                    try:
                        image_paragraph = document.add_paragraph()
                        image_paragraph.paragraph_format.space_before = Pt(4)
                        image_paragraph.paragraph_format.space_after = Pt(6)
                        image_run = image_paragraph.add_run()
                        image_run.add_picture(absolute_path, width=Inches(5.5))
                        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    except Exception:
                        note_paragraph = document.add_paragraph()
                        note_paragraph.add_run(
                            f"  [screenshot: {os.path.basename(absolute_path)}]"
                        ).font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                else:
                    note_paragraph = document.add_paragraph()
                    note_paragraph.add_run(
                        f"  [screenshot not found: {log.screenshot_path}]"
                    ).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

            self._add_horizontal_rule(document)

        footer_paragraph = document.add_paragraph()
        footer_paragraph.paragraph_format.space_before = Pt(18)
        footer_run = footer_paragraph.add_run(
            f"Generated by IBMTaskWeaver  •  {datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
